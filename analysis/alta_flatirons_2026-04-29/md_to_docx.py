"""
Convert the Alta Flatirons GC-facing markdown report to a structured .docx.
Handles: H1/H2/H3, bullet lists, tables (pipe-delimited), bold (**...**),
italic (_..._), inline code (`...`), and warning ⚠ glyph passthrough.

Run: .venv/bin/python analysis/alta_flatirons_2026-04-29/md_to_docx.py
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
SRC = HERE / 'alta_flatirons_480_report.md'
DST = HERE / 'alta_flatirons_480_report.docx'

INLINE_RE = re.compile(r'(\*\*[^*]+\*\*|_[^_]+_|`[^`]+`)')


def add_runs(paragraph, text):
    """Add text to a paragraph, parsing **bold**, _italic_, `code`."""
    for piece in INLINE_RE.split(text):
        if not piece:
            continue
        if piece.startswith('**') and piece.endswith('**'):
            run = paragraph.add_run(piece[2:-2])
            run.bold = True
        elif piece.startswith('_') and piece.endswith('_'):
            run = paragraph.add_run(piece[1:-1])
            run.italic = True
        elif piece.startswith('`') and piece.endswith('`'):
            run = paragraph.add_run(piece[1:-1])
            run.font.name = 'Menlo'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(piece)


def shade_cell(cell, color_hex):
    """Apply a fill color to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def parse_table(lines, start):
    """Read a markdown pipe table starting at lines[start]. Return (rows, next_index)."""
    rows = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith('|'):
        line = lines[i].strip()
        # Skip the separator row (|---|---|)
        if re.fullmatch(r'\|[\s\-:|]+\|', line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Light Grid Accent 1'
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # First row is header
            p = cell.paragraphs[0]
            p.text = ''
            add_runs(p, val)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
                shade_cell(cell, 'D9E2F3')


def main():
    md = SRC.read_text()
    lines = md.split('\n')

    doc = Document()
    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    # Default body font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            # Add a thin paragraph break with a bottom border via a thin rule
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '999999')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Headings
        if stripped.startswith('# '):
            p = doc.add_heading(level=0)
            add_runs(p, stripped[2:])
            i += 1
            continue
        if stripped.startswith('## '):
            p = doc.add_heading(level=1)
            add_runs(p, stripped[3:])
            i += 1
            continue
        if stripped.startswith('### '):
            p = doc.add_heading(level=2)
            add_runs(p, stripped[4:])
            i += 1
            continue

        # Table
        if stripped.startswith('|'):
            rows, next_i = parse_table(lines, i)
            add_table(doc, rows)
            i = next_i
            continue

        # Bullet
        if stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, stripped[2:])
            i += 1
            continue

        # Plain paragraph (may span multiple consecutive non-empty lines, but
        # we treat each line as its own paragraph for simplicity)
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(DST)
    print(f'Wrote: {DST}')


if __name__ == '__main__':
    main()
